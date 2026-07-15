"""
core/book_barcode_generator.py — Generator label QR Code buku massal

Alur:
1. Ambil data buku dari SQLite (semua / filter kategori / pilihan manual)
2. Generate gambar QR Code PNG per buku → folder book_barcode_images/
3. Generate .docx berisi tabel label QR Code (N label per baris) siap cetak

Format label per buku:
    ┌────────────────────────┐
    │     [ QR CODE ]        │
    │  BK-2025-0001          │
    │  Judul Buku            │
    │  Pengarang             │  ← opsional
    └────────────────────────┘

QR Code dipilih untuk label buku karena jauh lebih toleran terhadap
sudut/pencahayaan saat dibaca lewat webcam dibanding barcode linear, dan
alat pemindai USB 2D kini cukup terjangkau. Selaras dengan kartu anggota
(core/card_generator.py) dan filter scanner di config.py
(BARCODE_TYPE_FILTER = "QRCODE").
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from config import BASE_DIR
from core.label_utils import set_cell_border, set_cell_bg, set_cell_margin, make_qr_png

logger = logging.getLogger(__name__)

# ── Direktori output default ───────────────────────────────────────────────────
BOOK_BARCODE_DIR: Path = BASE_DIR / "assets" / "book_barcode_images"
BOOK_LABEL_DIR: Path   = BASE_DIR / "assets" / "book_label_output"

# Ukuran label buku dalam cm
LABEL_WIDTH_CM   = 6.0
LABEL_HEIGHT_CM  = 3.5
LABELS_PER_ROW   = 3   # 3 label per baris di halaman A4


# ══════════════════════════════════════════════════════════════════════════════
# Fungsi utilitas
# ══════════════════════════════════════════════════════════════════════════════

def _ensure_dirs(barcode_dir: Path, label_dir: Path) -> None:
    barcode_dir.mkdir(parents=True, exist_ok=True)
    label_dir.mkdir(parents=True, exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# Step 1: Ambil data buku dari database
# ══════════════════════════════════════════════════════════════════════════════

def fetch_books(
    kategori: Optional[str] = None,
    kode_list: Optional[list[str]] = None,
) -> list[dict]:
    """
    Ambil data buku dari SQLite untuk di-generate labelnya.

    Args:
        kategori  : filter per kategori (None = semua buku)
        kode_list : filter daftar kode_buku spesifik (None = tidak filter)

    Returns:
        list[dict] dengan key: kode_buku, judul, pengarang, kategori
    """
    from database.book_db import _get_connection
    conn = _get_connection()
    try:
        conditions = []
        params     = []

        if kategori and kategori != "Semua":
            conditions.append("kategori = ?")
            params.append(kategori)

        if kode_list:
            placeholders = ",".join("?" * len(kode_list))
            conditions.append(f"kode_buku IN ({placeholders})")
            params.extend(kode_list)

        where = f"WHERE {' AND '.join(conditions)}" if conditions else ""
        sql   = f"SELECT kode_buku, judul, pengarang, kategori FROM buku {where} ORDER BY kode_buku"

        rows = conn.execute(sql, params).fetchall()
        books = [
            {
                "kode_buku":  row[0],
                "judul":      row[1] or "",
                "pengarang":  row[2] or "",
                "kategori":   row[3] or "",
            }
            for row in rows
        ]
        logger.info("Fetch buku: %d buku ditemukan (kategori=%s)", len(books), kategori or "semua")
        return books
    finally:
        conn.close()


def fetch_all_categories() -> list[str]:
    """Ambil daftar kategori unik dari tabel buku."""
    from database.book_db import _get_connection
    conn = _get_connection()
    try:
        rows = conn.execute(
            "SELECT DISTINCT kategori FROM buku WHERE kategori IS NOT NULL AND kategori != '' ORDER BY kategori"
        ).fetchall()
        return [r[0] for r in rows]
    finally:
        conn.close()


# ══════════════════════════════════════════════════════════════════════════════
# Step 2: Generate gambar barcode PNG
# ══════════════════════════════════════════════════════════════════════════════

def generate_book_barcode_images(
    books: list[dict],
    barcode_dir: Optional[Path] = None,
    on_progress=None,
) -> dict[str, Path]:
    """
    Generate gambar QR Code PNG untuk setiap buku.
    (Parameter/nama fungsi tetap pakai awalan "barcode_" untuk konsistensi
    dengan core/card_generator.py.)

    Args:
        books       : list hasil fetch_books()
        barcode_dir : folder output PNG (default: BOOK_BARCODE_DIR)
        on_progress : callback(current, total)

    Returns:
        dict kode_buku → Path file PNG
    """
    if barcode_dir is None:
        barcode_dir = BOOK_BARCODE_DIR

    _ensure_dirs(barcode_dir, BOOK_LABEL_DIR)
    results: dict[str, Path] = {}
    total = len(books)

    for idx, book in enumerate(books):
        kode = book["kode_buku"]
        safe = kode.replace("/", "-").replace("\\", "-")
        out_path = barcode_dir / f"{safe}.png"

        try:
            make_qr_png(kode, out_path)
            results[kode] = out_path
            logger.debug("QR Code buku: %s → %s", kode, out_path.name)

        except Exception as exc:
            logger.error("Gagal generate QR Code buku %s: %s", kode, exc)

        if on_progress:
            on_progress(idx + 1, total)

    logger.info("QR Code buku selesai: %d/%d berhasil", len(results), total)
    return results


# ══════════════════════════════════════════════════════════════════════════════
# Step 3: Generate .docx label buku massal
# ══════════════════════════════════════════════════════════════════════════════

def _add_book_label_to_cell(
    cell,
    book: dict,
    barcode_path: Optional[Path],
    show_author: bool = True,
) -> None:
    """
    Isi cell dengan label buku:
        ┌──────────────────────┐
        │     [ QR CODE ]      │
        │  BK-2025-0001        │
        │  Judul Buku...       │
        │  Pengarang           │  ← jika show_author=True
        └──────────────────────┘
    """
    for p in cell.paragraphs:
        p.clear()

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell, color="999999", sz="6")
    set_cell_bg(cell, "FFFFFF")
    set_cell_margin(cell, top=60, bottom=60, left=80, right=80)

    # ── Gambar QR Code ────────────────────────────────────────────────────────
    p_bc = cell.paragraphs[0]
    p_bc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bc.paragraph_format.space_before = Pt(0)
    p_bc.paragraph_format.space_after  = Pt(2)

    if barcode_path and barcode_path.exists():
        run_bc = p_bc.add_run()
        # QR persegi → lebih kecil dari lebar wide-barcode lama (4.5cm) agar
        # muat proporsional dalam tinggi label 3.5cm bersama teks di bawahnya.
        run_bc.add_picture(str(barcode_path), width=Cm(2.6))
    else:
        r_na = p_bc.add_run("[no QR]")
        r_na.font.size = Pt(7)
        r_na.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Kode buku ─────────────────────────────────────────────────────────────
    p_kode = cell.add_paragraph()
    p_kode.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kode.paragraph_format.space_before = Pt(0)
    p_kode.paragraph_format.space_after  = Pt(1)
    r_kode = p_kode.add_run(book.get("kode_buku", ""))
    r_kode.bold = True
    r_kode.font.size = Pt(8)
    r_kode.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # ── Judul buku (potong jika terlalu panjang) ──────────────────────────────
    judul = book.get("judul", "")
    if len(judul) > 40:
        judul = judul[:38] + "…"

    p_judul = cell.add_paragraph()
    p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_judul.paragraph_format.space_before = Pt(0)
    p_judul.paragraph_format.space_after  = Pt(1)
    r_judul = p_judul.add_run(judul)
    r_judul.font.size = Pt(7.5)
    r_judul.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Pengarang (opsional) ──────────────────────────────────────────────────
    if show_author:
        pengarang = book.get("pengarang", "").strip()
        if pengarang and pengarang.lower() != "none":
            if len(pengarang) > 35:
                pengarang = pengarang[:33] + "…"
            p_pg = cell.add_paragraph()
            p_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_pg.paragraph_format.space_before = Pt(0)
            p_pg.paragraph_format.space_after  = Pt(0)
            r_pg = p_pg.add_run(pengarang)
            r_pg.italic    = True
            r_pg.font.size = Pt(7)
            r_pg.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def generate_book_labels_docx(
    books: list[dict],
    barcode_paths: dict[str, Path],
    labels_per_row: int = LABELS_PER_ROW,
    show_author: bool = True,
    label_dir: Optional[Path] = None,
    on_progress=None,
) -> Path:
    """
    Generate file .docx berisi label barcode buku massal.

    Args:
        books          : list buku dari fetch_books()
        barcode_paths  : dict dari generate_book_barcode_images()
        labels_per_row : jumlah label per baris (default 3)
        show_author    : tampilkan nama pengarang di label
        label_dir      : folder output .docx (default: BOOK_LABEL_DIR)
        on_progress    : callback(current, total)

    Returns:
        Path file .docx yang dihasilkan
    """
    if label_dir is None:
        label_dir = BOOK_LABEL_DIR

    _ensure_dirs(BOOK_BARCODE_DIR, label_dir)

    doc = Document()

    # ── Setup halaman A4 portrait ─────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(1.2)
    section.right_margin  = Cm(1.2)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    total  = len(books)
    n_rows = (total + labels_per_row - 1) // labels_per_row

    # ── Buat tabel grid label ─────────────────────────────────────────────────
    table = doc.add_table(rows=n_rows, cols=labels_per_row)
    table.style = "Table Grid"

    col_width = Cm(LABEL_WIDTH_CM)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # ── Isi setiap cell dengan label ──────────────────────────────────────────
    for idx, book in enumerate(books):
        row_idx = idx // labels_per_row
        col_idx = idx  % labels_per_row
        cell    = table.cell(row_idx, col_idx)
        kode    = book.get("kode_buku", "")
        bpath   = barcode_paths.get(kode)
        _add_book_label_to_cell(cell, book, bpath, show_author=show_author)

        if on_progress:
            on_progress(idx + 1, total)

    # Set tinggi baris
    for row in table.rows:
        row.height = Cm(LABEL_HEIGHT_CM)

    # ── Simpan ────────────────────────────────────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path  = label_dir / f"label_buku_{timestamp}.docx"
    doc.save(str(out_path))
    logger.info("Dokumen label buku berhasil: %s", out_path.name)
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# Entry point dari GUI
# ══════════════════════════════════════════════════════════════════════════════

def run_book_barcode_generation(
    kategori: Optional[str] = None,
    kode_list: Optional[list[str]] = None,
    show_author: bool = True,
    labels_per_row: int = LABELS_PER_ROW,
    output_dir: Optional[str | Path] = None,
    on_progress=None,
) -> dict:
    """
    Jalankan seluruh pipeline generate label barcode buku.

    Args:
        kategori       : filter kategori (None = semua)
        kode_list      : filter kode buku spesifik (None = tidak filter)
        show_author    : tampilkan pengarang di label
        labels_per_row : jumlah label per baris
        output_dir     : folder output kustom
        on_progress    : callback(stage, current, total)

    Returns dict:
        'books'       : list[dict] buku yang diproses
        'docx_path'   : Path file .docx
        'barcode_dir' : Path folder PNG
        'count'       : int jumlah label dibuat
        'errors'      : list[str]
    """
    if output_dir is None:
        barcode_dir = BOOK_BARCODE_DIR
        label_dir   = BOOK_LABEL_DIR
    else:
        output_dir  = Path(output_dir)
        barcode_dir = output_dir / "book_barcode_images"
        label_dir   = output_dir / "book_label_output"

    # Step 1: Fetch buku
    books = fetch_books(kategori=kategori, kode_list=kode_list)
    if not books:
        return {
            "books": [], "count": 0,
            "errors": ["Tidak ada buku yang ditemukan dengan filter tersebut."],
        }

    # Step 2: Generate barcode PNG
    def prog_barcode(cur, tot):
        if on_progress:
            on_progress("barcode", cur, tot)

    barcode_paths = generate_book_barcode_images(
        books, barcode_dir=barcode_dir, on_progress=prog_barcode
    )

    # Step 3: Generate .docx
    def prog_docx(cur, tot):
        if on_progress:
            on_progress("docx", cur, tot)

    docx_path = generate_book_labels_docx(
        books,
        barcode_paths,
        labels_per_row=labels_per_row,
        show_author=show_author,
        label_dir=label_dir,
        on_progress=prog_docx,
    )

    return {
        "books":       books,
        "docx_path":   docx_path,
        "barcode_dir": barcode_dir,
        "count":       len(books),
        "errors":      [],
    }