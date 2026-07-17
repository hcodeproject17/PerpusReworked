"""
core/card_generator.py — Generator kartu perpustakaan (QR Code) untuk
anggota yang sudah terdaftar di anggota.xlsx.

Alur:
1. Ambil anggota dari database/excel_reader.py (semua / filter / pilihan
   manual dari tab Anggota — lihat gui/card_tab.py)
2. Generate gambar QR Code per anggota → folder barcode_images/
3. Generate .docx berisi tabel kartu (N kartu per halaman) siap cetak

Pendaftaran anggota baru (termasuk assign ID Barcode & import massal
Excel) ada di database/excel_reader.py, dipakai dari gui/member_tab.py.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import BASE_DIR, EXCEL_COL_BARCODE, EXCEL_COL_NAME
from core.label_utils import set_cell_border, set_cell_bg, set_cell_margin, make_qr_png

logger = logging.getLogger(__name__)

# ── Direktori output ──────────────────────────────────────────────────────────
BARCODE_DIR: Path  = BASE_DIR / "assets" / "barcode_images"
CARDS_DIR: Path    = BASE_DIR / "assets" / "kartu_output"

# Ukuran kartu perpustakaan (kertas jeruk) dalam cm
CARD_WIDTH_CM  = 8.5
CARD_HEIGHT_CM = 5.5
CARDS_PER_ROW  = 2   # 2 kartu per baris di halaman A4


# ══════════════════════════════════════════════════════════════════════════════
# Fungsi utilitas
# ══════════════════════════════════════════════════════════════════════════════

def _ensure_dirs(barcode_dir: Path, cards_dir: Path) -> None:
    barcode_dir.mkdir(parents=True, exist_ok=True)
    cards_dir.mkdir(parents=True, exist_ok=True)


def fetch_members(query: Optional[str] = None) -> list[dict]:
    """Ambil data anggota dari anggota.xlsx untuk dipilih & dicetak kartunya."""
    from database.excel_reader import load_members, search as search_members
    try:
        return search_members(query) if query else load_members()
    except FileNotFoundError:
        return []


# ══════════════════════════════════════════════════════════════════════════════
# Step 1: Generate gambar QR Code
# ══════════════════════════════════════════════════════════════════════════════

def generate_barcode_images(
    members_with_id: list[dict],
    barcode_dir: Optional[Path] = None,
    on_progress=None,
) -> dict[str, Path]:
    """
    Generate gambar QR Code PNG untuk setiap anggota.
    (Fungsi/parameter tetap pakai awalan "barcode_" untuk konsistensi
    penamaan di modul ini.)

    Args:
        members_with_id : list anggota, tiap dict wajib punya EXCEL_COL_BARCODE
        barcode_dir     : folder output untuk gambar QR (default: BARCODE_DIR)
        on_progress     : callback(current, total) untuk progress bar GUI

    Returns:
        dict barcode_id → Path file PNG
    """
    if barcode_dir is None:
        barcode_dir = BARCODE_DIR

    _ensure_dirs(barcode_dir, CARDS_DIR)
    results: dict[str, Path] = {}
    total = len(members_with_id)

    for idx, member in enumerate(members_with_id):
        barcode_id = member[EXCEL_COL_BARCODE]
        safe_name  = barcode_id.replace("/", "-").replace("\\", "-")
        out_path   = barcode_dir / f"{safe_name}.png"

        try:
            make_qr_png(barcode_id, out_path)
            results[barcode_id] = out_path
            logger.debug("QR Code PNG: %s → %s", barcode_id, out_path.name)

        except Exception as exc:
            logger.error("Gagal generate QR Code %s: %s", barcode_id, exc)

        if on_progress:
            on_progress(idx + 1, total)

    logger.info("Generate QR Code selesai: %d/%d berhasil", len(results), total)
    return results


# ══════════════════════════════════════════════════════════════════════════════
# Step 2: Generate .docx kartu massal
# ══════════════════════════════════════════════════════════════════════════════

def _add_card_to_cell(
    outer_cell,
    member: dict,
    barcode_path: Optional[Path],
    school_name: str,
) -> None:
    """
    Isi outer_cell dengan kartu perpustakaan layout:

        ┌─────────────────────────────────────┐
        │   Nama Sekolah (bold, center)        │  ← header row, full width
        │   Kartu Perpustakaan (center)        │
        ├──────────────────────┬──────────────┤
        │  Nama: ...           │  [QR CODE]   │  ← info kiri, QR kanan
        │  ID Kartu: ...       │              │
        └──────────────────────┴──────────────┘
    """
    # Bersihkan isi default outer_cell
    for p in outer_cell.paragraphs:
        p.clear()
    outer_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Border luar kartu (tebal)
    set_cell_border(outer_cell, color="888888", sz="8")
    set_cell_margin(outer_cell, top=0, bottom=0, left=0, right=0)

    # ── Buat nested table 2 baris × 1 kolom di dalam outer_cell ──────────────
    # Baris 1: header (full width)
    # Baris 2: inner table 1×2 (info | barcode)
    p = outer_cell.paragraphs[0]._p

    nested = outer_cell.add_table(rows=2, cols=1)
    nested.style = "Table Grid"

    # Paksa nested table ikut lebar outer cell
    tbl  = nested._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")
    tblW.set(qn("w:type"), "pct")   # 100% lebar parent cell
    tblPr.append(tblW)

    # ── Baris 1: Header ───────────────────────────────────────────────────────
    header_cell = nested.cell(0, 0)
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(header_cell, "FFFFFF")
    set_cell_margin(header_cell, top=80, bottom=40, left=100, right=100)
    set_cell_border(header_cell, sides=["bottom"], color="888888", sz="6")

    # Kosongkan paragraf default
    for p in header_cell.paragraphs:
        p.clear()

    # Nama sekolah (bold, besar, center)
    p1 = header_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(2)
    r1 = p1.add_run(school_name)
    r1.bold      = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    # "Kartu Perpustakaan" (tidak bold, sedikit lebih kecil)
    p2 = header_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run("Kartu Perpustakaan")
    r2.bold      = False
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Baris 2: Info + Barcode ───────────────────────────────────────────────
    body_cell = nested.cell(1, 0)
    set_cell_bg(body_cell, "FFFFFF")
    set_cell_margin(body_cell, top=0, bottom=0, left=0, right=0)

    # Inner table 1×2 di dalam body_cell
    inner = body_cell.add_table(rows=1, cols=2)
    inner.style = "Table Grid"

    # Lebar: kiri 50%, kanan 50% (kolom barcode diperbesar agar lebih mudah discan)
    inner_tbl  = inner._tbl
    inner_tblPr = inner_tbl.find(qn("w:tblPr"))
    if inner_tblPr is None:
        inner_tblPr = OxmlElement("w:tblPr")
        inner_tbl.insert(0, inner_tblPr)
    inner_tblW = OxmlElement("w:tblW")
    inner_tblW.set(qn("w:w"),    "5000")
    inner_tblW.set(qn("w:type"), "pct")
    inner_tblPr.append(inner_tblW)

    # Set lebar kolom inner (dalam pct twips)
    for col_idx, col in enumerate(inner.columns):
        for c in col.cells:
            tcPr = c._tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            pct  = "2500" if col_idx == 0 else "2500"   # 50% : 50%
            tcW.set(qn("w:w"),    pct)
            tcW.set(qn("w:type"), "pct")
            tcPr.append(tcW)

    # ── Sel kiri: nama + id ───────────────────────────────────────────────────
    left_cell = inner.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margin(left_cell, top=80, bottom=80, left=100, right=60)
    set_cell_border(left_cell, sides=["right"], color="AAAAAA", sz="4")

    for p in left_cell.paragraphs:
        p.clear()

    nama = member.get(EXCEL_COL_NAME, "")
    bid  = member.get(EXCEL_COL_BARCODE, "")

    p_nama = left_cell.paragraphs[0]
    p_nama.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_nama.paragraph_format.space_before = Pt(0)
    p_nama.paragraph_format.space_after  = Pt(4)
    r_nama = p_nama.add_run(f"Nama: {nama}")
    r_nama.bold      = True
    r_nama.font.size = Pt(10)
    r_nama.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    p_id = left_cell.add_paragraph()
    p_id.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_id.paragraph_format.space_before = Pt(0)
    p_id.paragraph_format.space_after  = Pt(0)
    r_id = p_id.add_run(f"ID Kartu: {bid}")
    r_id.bold      = False
    r_id.font.size = Pt(9)
    r_id.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Kolom opsional tambahan (Kelas, NIS, dll)
    for key in member:
        if key not in (EXCEL_COL_BARCODE, EXCEL_COL_NAME):
            val = str(member.get(key, "") or "").strip()
            if val and val.lower() != "none":
                px = left_cell.add_paragraph()
                px.alignment = WD_ALIGN_PARAGRAPH.LEFT
                px.paragraph_format.space_before = Pt(0)
                px.paragraph_format.space_after  = Pt(0)
                rx = px.add_run(f"{key}: {val}")
                rx.font.size = Pt(8)
                rx.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Sel kanan: gambar QR Code ─────────────────────────────────────────────
    right_cell = inner.cell(0, 1)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margin(right_cell, top=60, bottom=60, left=60, right=60)

    for p in right_cell.paragraphs:
        p.clear()

    if barcode_path and barcode_path.exists():
        p_bc = right_cell.paragraphs[0]
        p_bc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_bc.paragraph_format.space_before = Pt(0)
        p_bc.paragraph_format.space_after  = Pt(0)
        run_bc = p_bc.add_run()
        run_bc.add_picture(str(barcode_path), width=Cm(3.3))  # QR persegi, disesuaikan agar muat tinggi kartu
    else:
        p_bc = right_cell.paragraphs[0]
        p_bc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_na = p_bc.add_run("[no QR]")
        r_na.font.size = Pt(7)
        r_na.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def generate_cards_docx(
    members_with_id: list[dict],
    barcode_paths: dict[str, Path],
    school_name: str = "SMA Negeri 1",
    cards_per_row: int = CARDS_PER_ROW,
    cards_dir: Optional[Path] = None,
    on_progress=None,
) -> Path:
    """
    Generate file .docx berisi kartu perpustakaan massal.
    Setiap halaman berisi grid kartu dengan tata letak tabel.

    Args:
        members_with_id : list anggota dengan barcode ID
        barcode_paths   : dict dari generate_barcode_images()
        school_name     : nama sekolah untuk header kartu
        cards_per_row   : jumlah kartu per baris (default 2)
        cards_dir       : folder output untuk file .docx (default: CARDS_DIR)
        on_progress     : callback(current, total)

    Returns:
        Path file .docx yang dihasilkan
    """
    if cards_dir is None:
        cards_dir = CARDS_DIR

    _ensure_dirs(BARCODE_DIR, cards_dir)

    doc = Document()

    # ── Setup halaman A4 portrait ────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── Hitung baris tabel ────────────────────────────────────────────────────
    total = len(members_with_id)
    n_rows = (total + cards_per_row - 1) // cards_per_row

    # ── Buat tabel ────────────────────────────────────────────────────────────
    table = doc.add_table(rows=n_rows, cols=cards_per_row)
    table.style = "Table Grid"

    # Set lebar kolom
    col_width = Cm(CARD_WIDTH_CM)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # ── Isi tabel dengan kartu ────────────────────────────────────────────────
    for idx, member in enumerate(members_with_id):
        row_idx = idx // cards_per_row
        col_idx = idx  % cards_per_row
        cell    = table.cell(row_idx, col_idx)
        barcode_id   = member.get(EXCEL_COL_BARCODE, "")
        barcode_path = barcode_paths.get(barcode_id)
        _add_card_to_cell(cell, member, barcode_path, school_name)

        if on_progress:
            on_progress(idx + 1, total)

    # Set tinggi baris
    for row in table.rows:
        row.height = Cm(CARD_HEIGHT_CM)

    # ── Simpan ────────────────────────────────────────────────────────────────
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path  = cards_dir / f"kartu_perpustakaan_{timestamp}.docx"
    doc.save(str(out_path))
    logger.info("Dokumen kartu berhasil dibuat: %s", out_path.name)
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# Fungsi utama (entry point dari GUI)
# ══════════════════════════════════════════════════════════════════════════════

def run_card_generation(
    members: list[dict],
    school_name: str,
    cards_per_row: int = CARDS_PER_ROW,
    output_dir: Optional[str | Path] = None,
    on_progress=None,
) -> dict:
    """
    Cetak kartu QR Code untuk anggota yang sudah terdaftar (dipilih dari
    tab Anggota — lihat gui/card_tab.py). Tidak menulis apa pun ke
    anggota.xlsx; pendaftaran anggota baru ada di gui/member_tab.py.

    Args:
        members       : list anggota (dari core.card_generator.fetch_members())
        school_name   : nama sekolah untuk header kartu
        cards_per_row : jumlah kartu per baris (default 2)
        output_dir    : folder output kustom untuk QR dan docx (default: assets/)
        on_progress   : callback(stage, current, total)

    Returns dict dengan key:
        'members'     : list[dict] anggota yang diproses
        'docx_path'   : Path file .docx output
        'barcode_dir' : Path folder gambar QR Code
        'count'       : int jumlah kartu dibuat
        'errors'      : list[str] error fatal
    """
    if output_dir is None:
        barcode_dir = BARCODE_DIR
        cards_dir = CARDS_DIR
    else:
        output_dir = Path(output_dir)
        barcode_dir = output_dir / "barcode_images"
        cards_dir = output_dir / "kartu_output"

    if not members:
        return {"members": [], "count": 0, "errors": ["Tidak ada anggota yang dipilih."]}

    def prog_barcode(cur, tot):
        if on_progress:
            on_progress("barcode", cur, tot)

    barcode_paths = generate_barcode_images(members, barcode_dir=barcode_dir, on_progress=prog_barcode)

    def prog_docx(cur, tot):
        if on_progress:
            on_progress("docx", cur, tot)

    docx_path = generate_cards_docx(
        members,
        barcode_paths,
        school_name=school_name,
        cards_per_row=cards_per_row,
        cards_dir=cards_dir,
        on_progress=prog_docx,
    )

    return {
        "members":     members,
        "docx_path":   docx_path,
        "barcode_dir": barcode_dir,
        "count":       len(members),
        "errors":      [],
    }